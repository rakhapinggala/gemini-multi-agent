from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os

def clean_markdown(text):
    """Hapus sintaks markdown sederhana seperti **bold**."""
    return text.replace("**", "").replace("*", "")

def generate_report(query, answer, category_label, sources, filename=None):
    if not filename:
        filename = f"static/reports/laporan_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"

    doc = Document()

    # Title
    title = doc.add_heading('LAPORAN CHATBOT', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Metadata
    doc.add_paragraph().add_run("Tanggal     : ").bold = True
    doc.paragraphs[-1].add_run(datetime.now().strftime('%Y-%m-%d %H:%M'))

    doc.add_paragraph().add_run("Kategori    : ").bold = True
    doc.paragraphs[-1].add_run(category_label)

    # Pertanyaan
    doc.add_heading("Pertanyaan", level=1)
    doc.add_paragraph(clean_markdown(query))

    # Jawaban
    doc.add_heading("Jawaban", level=1)
    doc.add_paragraph(clean_markdown(answer))

    # Sumber
    if sources:
        doc.add_heading("Sumber Informasi", level=1)
        for doc_source in sources:
            name = doc_source.metadata.get("source", "Tidak diketahui")
            excerpt = clean_markdown(doc_source.page_content[:300].strip().replace("\n", " "))
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(f"{name} ").bold = True
            para.add_run(f"{excerpt}")

    # Format default
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    os.makedirs(os.path.dirname(filename), exist_ok=True)
    doc.save(filename)

    return filename.replace("static/", "")
